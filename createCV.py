from docx import Document 
from docx.shared import Cm

#Variables
document = Document()
picture_route = input("Insert the absolute route of your cv picture: ")
name = input("Write your name and lastname: ")
phone_number = input("What is your phone number: ")
email = input("Now, write your email: ")
about_me = input("Tell me about yourself ")

#Execution code
document.add_picture(
    picture_route, 
    width=Cm(2.0)
    )
document.add_paragraph(f"{name}  |  {phone_number}  | {email}")
#Add About Me
document.add_heading("About me")
document.add_paragraph(about_me)

# Work Experience
print("Now we are going to talk about your working experience.")
document.add_heading("Work Experience")
p = document.add_paragraph()


confirm = " "
while confirm != "no":
    company = input("Enter the name of the company: ")
    from_date = input("Write your start date: ")
    to_date = input("Write your end date: ")
    exp_description = input(f"Describe your experience at {company}: ")
    p.add_run(f"{company}  ").bold = True
    p.add_run(f"{from_date} - {to_date} \n").italic = True
    p.add_run(exp_description)
    confirm = input("Do you want to add more working experience(YES/no)? ")


# Save the document
document.save("cv.docx")
